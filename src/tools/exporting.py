import os

from docx import Document
from win32com import client

from models.account import Account
from models.daybook import Daybook
from models.donor import Donor
from models.record import Record


def _add_heading(doc: Document, text: str):
    doc.add_heading(text)


def _add_paragraph(doc: Document, text: str):
    paragraph = doc.add_paragraph()
    paragraph.add_run(text)


def _add_records_table(doc: Document, records: [Record]):
    table = doc.add_table(rows=len(records) + 1, cols=6)

    headers = ["#", "Belegdatum", "Buchungstext", "Betrag", "Soll-Konto", "Haben-Konto"]
    for col, header in enumerate(headers):
        table.cell(0, col).text = header

    for row, record in enumerate(records, start=1):
        cells = table.rows[row].cells
        cells[0].text = str(record.number)
        cells[1].text = str(record.date.strftime("%d.%m.%Y"))
        cells[2].text = str(record.description)
        cells[3].text = str(record.amount).replace(".", ",")
        cells[4].text = str(record.debit_account)
        cells[5].text = str(record.credit_account)


def _save(doc: Document, file_name: str):
    doc.save(file_name + ".docx")

    doc_path = os.path.abspath(file_name + ".docx")
    pdf_path = os.path.abspath(file_name + ".pdf")

    word = client.Dispatch("Word.Application")
    doc = word.Documents.Open(doc_path)
    doc.SaveAs(pdf_path, FileFormat=17)
    doc.Close()
    word.Quit()


def _create_daybook_report(daybook: Daybook):
    doc = Document()
    _add_heading(doc, "Tagebuch")
    _add_records_table(doc, daybook.records)
    _save(doc, "Tagebuch")


def _create_donors_report(donors: [Donor]):
    doc = Document()
    file_name = "Spendenbericht"
    _add_heading(doc, file_name)
    _save(doc, file_name)


def _create_accounts_report(accounts: [Account]):
    doc = Document()
    file_name = "Konten"
    _add_heading(doc, file_name)
    _save(doc, file_name)


def _create_summery_report(accounts: [Account]):
    doc = Document()
    file_name = "Finanzbericht"
    _add_heading(doc, file_name)
    _save(doc, file_name)


def export(report):
    _create_accounts_report(report.accounts)
    _create_daybook_report(report.daybook)
    _create_donors_report(report.donors)
    _create_summery_report(report.accounts)
